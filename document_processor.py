# document_processor.py - Updated for Flask app
import os
import logging
from typing import List, Dict, Optional
from datetime import datetime
import hashlib
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import PyPDF2
import docx

logger = logging.getLogger(__name__)

class CompanyDocumentProcessor:
    """Handles document ingestion and chunking for company documents"""
    
    def __init__(self, chunk_size: int = 400, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
            length_function=len
        )
    
    def load_documents(self, file_paths: List[str]) -> List[Document]:
        """Load company documents from various file types"""
        documents = []
        
        for file_path in file_paths:
            try:
                # Determine file type and extract text
                file_extension = os.path.splitext(file_path)[1].lower()
                
                if file_extension == '.txt':
                    content = self._load_text_file(file_path)
                elif file_extension == '.pdf':
                    content = self._load_pdf_file(file_path)
                elif file_extension in ['.doc', '.docx']:
                    content = self._load_word_file(file_path)
                else:
                    logger.warning(f"Unsupported file type: {file_extension}")
                    continue
                
                # Extract metadata
                doc_type = self._classify_document(file_path, content)
                
                doc = Document(
                    page_content=content,
                    metadata={
                        "source": os.path.basename(file_path),
                        "doc_type": doc_type,
                        "last_updated": datetime.now().isoformat(),
                        "checksum": hashlib.md5(content.encode()).hexdigest()
                    }
                )
                documents.append(doc)
                logger.info(f"Loaded document: {os.path.basename(file_path)}")
                
            except Exception as e:
                logger.error(f"Error loading {file_path}: {str(e)}")
                
        return documents
    
    def _load_text_file(self, file_path: str) -> str:
        """Load text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _load_pdf_file(self, file_path: str) -> str:
        """Load PDF file"""
        text = ""
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
        return text
    
    def _load_word_file(self, file_path: str) -> str:
        """Load Word document"""
        doc = docx.Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text)
    
    def _classify_document(self, file_path: str, content: str) -> str:
        """Classify document type based on content and filename"""
        file_name = os.path.basename(file_path).lower()
        content_lower = content.lower()[:1000]  # Check first 1000 chars
        
        # Check filename patterns
        if any(term in file_name for term in ['faq', 'frequently']):
            return "faq"
        elif any(term in file_name for term in ['policy', 'compliance', 'kyc', 'aml']):
            return "policy"
        elif any(term in file_name for term in ['loan', 'mortgage', 'credit']):
            return "loan"
        elif any(term in file_name for term in ['account', 'savings', 'checking']):
            return "account"
        elif any(term in file_name for term in ['card', 'debit', 'credit']):
            return "card"
        
        # Check content patterns
        if any(term in content_lower for term in ['frequently asked', 'q:', 'question:']):
            return "faq"
        elif any(term in content_lower for term in ['compliance', 'regulation', 'requirement']):
            return "policy"
        elif any(term in content_lower for term in ['interest rate', 'apr', 'loan term']):
            return "loan"
        
        return "general"
    
    def chunk_documents(self, documents: List[Document]) -> List[Document]:
        """Split documents into chunks with metadata preservation"""
        all_chunks = []
        
        for doc in documents:
            chunks = self.text_splitter.split_text(doc.page_content)
            
            for i, chunk in enumerate(chunks):
                chunk_doc = Document(
                    page_content=chunk,
                    metadata={
                        **doc.metadata,
                        "chunk_id": f"{doc.metadata['source']}_{i}",
                        "chunk_index": i,
                        "total_chunks": len(chunks)
                    }
                )
                all_chunks.append(chunk_doc)
                
        logger.info(f"Created {len(all_chunks)} chunks from {len(documents)} documents")
        return all_chunks